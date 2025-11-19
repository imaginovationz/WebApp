import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import traceback
from playwright.sync_api import sync_playwright, TimeoutError as pwTimeoutError
from playwright.sync_api import Page
import openpyxl
from openpyxl.styles import Font, Alignment
from time import sleep
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from  ConfigFiles.CIBCDesktop.TestConfig import login_url, username, password, supportUser, supportPassword, glb_wait, headless_flag
import time
import pandas as pd

# Define all XPaths in one place to follow the Page Object Model (POM) pattern
XPATHS = {
    "deal_app_id": '//*[@id="deal_app_id"]',
    "applicant_name": '//*[@id="dealsummary_applicant_name"]',
    "deal_status": '//*[@id="deal_status_id"]',
    "product_selection": '//*[@id="productSelFA"]',
    "deal_purpose": '#B1_ProdDet_MortgageDetails_MortgagePurpose',
    "dfr": '#B1_ProdDet_MortgageDetails_DateFundsRequired',
    "amendment_history": '//*[@id="dAmendmentHistory"]',
    "collapse_one": '//*[@id="collapseOne"]/div/ul[6]',
    "documents": '//*[@id="documentsFA"]',
    "history": '//*[@id="dHistoryFA"]',
    "notifications": '//*[@id="dNotifFA"]',
    "mobile_notifications": '//*[@id="mobileNotif"]'
}


def create_word_document(deal,match_response,df):
    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading(f'RESL_SUMMARY_REPORT_{deal}', level=1)

    # Surround the heading '<Pre-Condition>' in a table
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = '<Pre-Condition>'
    doc.add_heading('MMTG RESL Report Data', level=2)
    
    if not df.empty:
        table = doc.add_table(rows=(len(df) + 1), cols=len(df.columns))
        table.style = 'Table Grid'

        for col_idx, column_name in enumerate(df.columns):
            table.cell(0, col_idx).text = column_name

        for row_idx, row in enumerate(df.itertuples(index=False), start=1):
            for col_idx, value in enumerate(row):
                table.cell(row_idx, col_idx).text = str(value)

    doc.add_heading('mMortgage #', level=2)
    doc.add_paragraph(match_response.get('mMortgage #', 'N/A'))
    doc.add_picture('mmtg_num.png', width=Inches(8))
    
    # Add the first set of fields and attach the screenshot
    doc.add_heading('CLASS Application #', level=2)
    doc.add_paragraph(match_response.get('CLASS Application #', 'N/A'))
    doc.add_heading('Primary Applicant Name', level=2)
    doc.add_paragraph(match_response.get('Primary Applicant Name', 'N/A'))
    doc.add_heading('Deal Status', level=2)
    doc.add_paragraph(match_response.get('Deal Status', 'N/A'))

    doc.add_picture('highlighted_deal_app_id.png', width=Inches(8))


    # Add the second set of fields and attach the screenshot
    doc.add_heading('Deal Purpose', level=2)
    doc.add_paragraph(match_response.get('Deal Purpose', 'N/A'))
    doc.add_picture('highlighted_prod_selection.png', width=Inches(8))

    doc.add_heading('DFR Less Than 3 Days', level=2)
    doc.add_paragraph(match_response.get('DFR Less Than 3 Days', 'N/A'))

    # Add the third set of fields and attach the screenshot
    doc.add_heading('Approval Date', level=2)
    doc.add_paragraph(match_response.get('Approval Date', 'N/A'))
    doc.add_picture('approval_date.png', width=Inches(8))

    # Add other fields and screenshots as per the code
    doc.add_heading('# of Internal Payout with a Funding Condition', level=2)
    doc.add_paragraph(match_response.get('# of Internal Payout with a Funding Condition of Payout close, Limit Decrease, Limit Decrease and Paydown', 'N/A'))
    doc.add_picture('liability.png', width=Inches(8))

    doc.add_heading('Payout LOD Form Initial Upload', level=2)
    doc.add_paragraph(match_response.get('Payout LOD Form Initial Upload', 'N/A'))
    doc.add_picture('lod_date.png', width=Inches(8))

    doc.add_heading('Payout LOD Form Recent Upload', level=2)
    doc.add_paragraph(match_response.get('Payout LOD Form Recent Upload', 'N/A'))
    doc.add_picture('lod_payout_date.png', width=Inches(8))



    # Before saving the Word document, delete it if it already exists
    file_path = f'RESL_SUMMARY_REPORT_{deal}.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Deleted existing file: {file_path}")

    doc.save(file_path)
    print(f"Word document RESL_SUMMARY_REPORT_{deal}.docx created successfully.")

def wait_and_click(page: Page, selector: str, expect_navigation: bool = False, timeout: int = glb_wait):
    """
    Waits for a selector and performs a click, with navigation and page load wait.

    :param page: Playwright Page object
    :param selector: CSS selector to click
    :param expect_navigation: Whether to expect navigation after the click
    :param timeout: Max wait time in milliseconds
    """
    page.wait_for_selector(selector, timeout=timeout)

    if expect_navigation:
        with page.expect_navigation(timeout=timeout):
            print(f"selector{selector} and timeout= {timeout}")
            page.click(selector)
    else:
        page.click(selector)
        print(f"Clicking selector {selector}")

    wait_for_page_load(page)

def wait_for_page_load(page: Page):
    page.wait_for_load_state("load", timeout=glb_wait)
    # page.wait_for_load_state("networkidle", timeout=glb_wait)
    page.wait_for_load_state("domcontentloaded")

def wait_and_click_dont_throw_exception_if_element_not_exsist(page: Page, selector: str, expect_navigation: bool = False, timeout: int = glb_wait):
    """
    Waits for a selector and performs a click, with navigation and page load wait.

    :param page: Playwright Page object
    :param selector: CSS selector to click
    :param timeout: Max wait time in milliseconds
    """
    wait_for_page_load(page)
    try:
        page.wait_for_selector(selector, timeout=timeout)
        element = page.query_selector(selector)
        if element is None:
                print(f"Element with selector '{selector}' not found. Continuing...")
                return False
        wait_and_click(page,selector,expect_navigation)
    except pwTimeoutError:
        return False
    
def click_if_exists(page: Page, selector: str):
    """
        default wait is only for 5 seconds
    """
    try:
        page.wait_for_selector(selector, timeout=5000)  
        page.click(selector)
        wait_for_page_load(page)
        print(f"Clicked element with selector: {selector}")
    except Exception as e:
        print(f"Element with selector: {selector} not found, continuing...")

# Open the Excel file and read the data into a pandas DataFrame
def read_excel_to_dataframe(file_path):
    try:
        df = pd.read_excel(file_path)
        print("Excel file successfully read into a DataFrame.")
        print(df)  # Display all rows of the DataFrame
        return df
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None
def delete_images():
    image_extensions = [".png", ".jpg", ".jpeg"]
    for file in os.listdir():
        if any(file.endswith(ext) for ext in image_extensions):
            os.remove(file)
            print(f"Deleted image: {file}")

def handle_angular_popups(page: Page):
    """
    Handle pop-up by closing x
    """
    try:
        # Wait briefly for popups to show

        # Common popup/modal containers in Angular apps
        form_containers = page.query_selector_all('.msgModalContent.popup-size')

        for container  in form_containers:
            cancel_btn = container.query_selector('.msgModalHeader > button.close')
            if cancel_btn:
                print(f"Popup detected")
                cancel_btn.click()
                print(f"Clicked popup button: {cancel_btn}")
                break

    except Exception as e:
        print(f"Error handling popups: {e}")

def wait_and_fill(page: Page, selector: str, value: str, timeout: int = glb_wait):
    """
    Waits for a selector and fills the input field.
    :param page: Playwright Page object
    :param selector: CSS selector for the input field
    :param value: Value to input
    :param timeout: Max wait time in milliseconds
    """
    wait_for_page_load(page)
    page.wait_for_selector(selector, timeout=timeout)
    page.fill(selector, value)

def handle_file_download(page, download_icon):
    """Handles the file download process for a given download icon."""
    if download_icon:
        try:
            # Use Playwright's event listener to handle download
            print("Setting up download handler...")
            with page.expect_download() as download_info:
                # Click the download button
                print("Clicking download icon...")
                download_icon.click(force=True)
                
            # Get the download object
            download = download_info.value
            original_filename = download.suggested_filename
            print(f"Download started: {original_filename}")
            
            # Wait for the download to complete
            print("Waiting for download to complete...")
            temp_path = download.path()
            print(f"Downloaded to temporary path: {temp_path}")
            
            # Create a 'downloads' folder in the script directory
            script_folder = os.path.dirname(os.path.abspath(__file__))
            downloads_folder = os.path.join(script_folder, "downloads")
            os.makedirs(downloads_folder, exist_ok=True)
            
            # Move the file to the 'downloads' folder with its original name
            final_path = os.path.join(downloads_folder, original_filename)
            import shutil
            shutil.move(temp_path, final_path)
            print(f"File moved to downloads folder: {final_path}")
            df = None
            import numpy as np
            # Attempt to read the Excel file
            try:
                df = pd.read_excel(final_path)
                df = df.fillna("").applymap(lambda x: "" if pd.isna(x) else x)
                print("Excel file successfully read. Preview:")
                print(df)  # Display all rows of the DataFrame
            except Exception as excel_err:
                print(f"Warning: Could not read Excel file: {excel_err}")
                
            return df
        except Exception as download_error:
            print(f"Error during download process: {download_error}")
            print(traceback.format_exc())

def get_deal_summary_data(page,glb_wait,df,deal):
    match_response = {}
    delete_images()
    print(df)
    rep_class_number = ""
    # Extract and strip the first non-NA value from the 'mMortgage #' column
    if "mMortgage #" in df.columns and not df["mMortgage #"].dropna().empty:
        rep_class_number = df["mMortgage #"].dropna().iloc[0]

    print(type(rep_class_number), rep_class_number)
    deal_app_id_element = page.query_selector(XPATHS["deal_app_id"])
    if deal_app_id_element:
        # Extract the text from the element
        deal_app_id_text = deal_app_id_element.inner_text().strip()
        print(f"Extracted text from deal_app_id: {deal_app_id_text}")

        #  this need to restore afte fix  if deal_app_id_text.replace('\u00A0', ' ').strip() == rep_class_number.replace('\u00A0', ' ').strip():
        if deal_app_id_text == rep_class_number:

            match_response['mMortgage #'] = "matched"
            page.evaluate("element => element.style.border = '2px solid green'", deal_app_id_element)
            print(f"deal_app_id: {deal_app_id_text} matches with mMortgage #")
        else:
            match_response['mMortgage #'] = "not matched"
            page.evaluate("element => element.style.border = '2px solid red'", deal_app_id_element)
            print(f"deal_app_id: {deal_app_id_text} not  with mMortgage #")
    

    deal_applicant_name = df["Primary Applicant Name"].dropna().iloc[0].strip()
    applicant_name = page.query_selector(XPATHS["applicant_name"])
    if applicant_name:
        # Extract the text from the element
        applicant_name_text = applicant_name.inner_text().strip()
        print(f"Extracted text from deal_app_id: {applicant_name_text} and applicant_name: {deal_applicant_name}")
        if deal_applicant_name.replace('\u00A0', ' ').strip() == applicant_name_text.replace('\u00A0', ' ').strip():
            match_response['Primary Applicant Name'] = "matched"
            page.evaluate("element => element.style.border = '2px solid green'", applicant_name)
            print(f"deal_applicant_name: {deal_applicant_name} matches with Primary Applicant Name")
        else:
            match_response['Primary Applicant Name'] = "not matched"
            page.evaluate("element => element.style.border = '2px solid red'", applicant_name)
            print(f"deal_applicant_name: {deal_applicant_name} not matches with Primary Applicant Name")

    df_deal_status = df["Deal Status"].dropna().iloc[0].strip()
    deal_status = page.query_selector(XPATHS["deal_status"])
    if deal_status:
        # Extract the text from the element
        deal_status_text = deal_status.inner_text().strip()
        print(f"Extracted text from deal_app_id: {deal_status} and deal_status: {df_deal_status}")
        if df_deal_status.replace('\u00A0', ' ').strip() == deal_status_text.replace('\u00A0', ' ').strip():
            match_response['Deal Status'] = "matched"
            page.evaluate("element => element.style.border = '2px solid green'", deal_status)
            print(f"Deal status: {df_deal_status} matches with Deal Status")
        else:
            match_response['Deal Status'] = "not matched"
            page.evaluate("element => element.style.border = '2px solid red'", deal_status)
            print(f"Deal status: {df_deal_status} not match with Deal Status")

    page.screenshot(path="highlighted_deal_app_id.png")
    
    wait_and_click(page,XPATHS["product_selection"])
    handle_angular_popups(page)
    page.wait_for_timeout(glb_wait/2)
    match_response['Deal Purpose'] = "not matched"
    df_deal_purpose = ""
    if "Deal Purpose" in df.columns and not df["Deal Purpose"].dropna().empty:
        df_deal_purpose = df["Deal Purpose"].dropna().iloc[0].strip()
    deal_purpose = page.query_selector(XPATHS["deal_purpose"])
    if deal_purpose:
        # Extract the text from the element
        selected_option = deal_purpose.query_selector('option:checked')
        deal_purpose_text = ""
        if selected_option:
            deal_purpose_text = selected_option.inner_text().strip()
        print(f"Extracted text from deal_app_id: {deal_purpose_text} and deal_purpose: {df_deal_purpose}")
        if df_deal_purpose.replace('\u00A0', ' ').strip() == deal_purpose_text.replace('\u00A0', ' ').strip():
            match_response['Deal Purpose'] = "matched"
            page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", deal_purpose)
        else:
            match_response['Deal Purpose'] = "not matched"
            page.evaluate("element => { element.style.border = '2px solid red'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", deal_purpose)
    if df_deal_purpose == "" and match_response['Deal Purpose'] == "not matched":  # require to handle where both are empt
        match_response['Deal Purpose'] = "matched"

    match_response['DFR Less Than 3 Days'] = "not matched" # explicitly set to not matched because validation will be in else statement
    df_dfr = df["Date Funds Required (DFR) date"].dropna().iloc[0].strip()
    dfr = page.query_selector(XPATHS["dfr"])
    if dfr:
        dfr.scroll_into_view_if_needed()
        dfr_text = page.evaluate("element => element.querySelector('input').value", dfr)
        print(f"Extracted text from dfr: {dfr_text} and dfr: {df_dfr}")
        if df_dfr.replace('\u00A0', ' ').strip() == dfr_text.replace('\u00A0', ' ').strip():
            match_response['Date Funds Required (DFR) date'] = "matched"
            page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", dfr)
            print(f"DFR: {df_dfr} matches") 
        else:
            match_response['Date Funds Required (DFR) date'] = "not matched"
            page.evaluate("element => { element.style.border = '2px solid red'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", dfr)
            print(f"DFR: {df_dfr} not matches") 


        # Parse the extracted date and check if it's less than 3 days from today
        if dfr_text:
            df_ind = df["DFR Date less then 3 days – (Y/N)"].dropna().iloc[0].strip()
            dfr_date = datetime.strptime(dfr_text, "%m/%d/%Y")
            today = datetime.now()
            if (dfr_date - today).days < 3:
                if df_ind == "Y":
                    match_response['DFR Less Than 3 Days'] = "matched"
            else:
                if df_ind == "N":
                    match_response['DFR Less Than 3 Days'] = "not matched"
            print(f"DFR date: {dfr_date}, Less than 3 days: {match_response['DFR Less Than 3 Days']}")

    page.screenshot(path="highlighted_prod_selection.png")


    wait_and_click(page,XPATHS["amendment_history"])
    handle_angular_popups(page)
    wait_for_page_load(page)
    df_dfr = ""
    if "Approval Date" in df.columns and not df["Approval Date"].dropna().empty:
        df_dfr = df["Approval Date"].dropna().iloc[0].strip()

    match_response['Approval Date'] = "not matched" # explicitly set to not matched because validation will be in else statement
    page.wait_for_timeout(glb_wait/2)
    rows = page.query_selector_all("div.mrow.col-md-12.px-0.row")
    for row in rows:
        columns = row.query_selector_all("div.col-md-2.pr-0")
        if len(columns) >= 6:
            second_col = columns[1].inner_text().strip()
            third_col = columns[2].inner_text().strip()
            fourth_col = columns[3].inner_text().strip()
            # print(f"Extracted text from columns: {second_col}, {third_col}, {fourth_col}")
            if second_col == "Submitted" and third_col == "Approved" and fourth_col == "Approved":
                timestamp = columns[5].inner_text().strip()
                match_response['Approval Date'] = "matched"
                page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", row)
                print(f"approval_date: {timestamp} matches") 
                break
    if df_dfr == "" and match_response['Approval Date'] == "not matched":  # require to handle where both are empt
        match_response['Approval Date'] = "matched"
    page.screenshot(path="approval_date.png")


    wait_and_click(page,XPATHS["collapse_one"])
    handle_angular_popups(page)
    wait_for_page_load(page)
    cibc_liability = df["# of Internal Payout with a Funding Condition of Payout close, Limit Decrease, Limit Decrease and Paydown"].dropna().iloc[0]
    match_response['# of Internal Payout with a Funding Condition of Payout close, Limit Decrease, Limit Decrease and Paydown'] = "not matched" # explicitly set to not matched because validation will be in else statement
    page.wait_for_timeout(glb_wait/2)
    rows = page.query_selector_all("div.row.py-3.mrow.ng-star-inserted")
    count = 0
    for row in rows:
        payout_icon = row.query_selector("img[src='./assets/images/Payout Icon.svg']")
        second_last_col = row.query_selector(f"div[id^='liab_TabelHed_creditor_name_']")  
        if payout_icon and second_last_col:
            second_last_col_text = second_last_col.inner_text().strip()
            page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", row)
            if second_last_col_text == "CIBC":
                count += 1
    print(f"Count of rows with Payout Icon and CIBC in the second last column: {count}")
    if cibc_liability == count:
        match_response['# of Internal Payout with a Funding Condition of Payout close, Limit Decrease, Limit Decrease and Paydown'] = "matched"
    page.screenshot(path="liability.png")


    wait_and_click(page,XPATHS["documents"])
    handle_angular_popups(page)
    wait_for_page_load(page)
    page.wait_for_timeout(glb_wait/2)
    # Safely handle empty or NaN values in the 'Payout LOD Form Initial Upload' column
    lod_date = ""
    if "Payout LOD Form Initial Upload" in df.columns and not df["Payout LOD Form Initial Upload"].dropna().empty:
        lod_date = df["Payout LOD Form Initial Upload"].dropna().iloc[0].strip()
    match_response['Payout LOD Form Initial Upload'] = "not matched"
    lod_date_element = page.query_selector('//label[@class = "doc_dispUploadDocname_Underline" and (contains(text(),"Letter of Direction") or contains(text(),"Lettre d’instructions"))]//following::label[@class="doc_label_teritory"]')
    if lod_date_element:
            # Extract the text from the element
            lod_date_text = lod_date_element.inner_text().strip()
            lod_date_element.scroll_into_view_if_needed()
            print(f"Extracted text from lod_date_text: {lod_date_text} and lod_date: {lod_date}")
            if lod_date_text.replace('\u00A0', ' ').strip() == lod_date.replace('\u00A0', ' ').strip():
                match_response['Payout LOD Form Initial Upload'] = "matched"
                page.evaluate("element => element.style.border = '2px solid green'", lod_date_element)
                print(f"lod_date: {lod_date} matches")
            else:
                match_response['Payout LOD Form Initial Upload'] = "not matched"
                page.evaluate("element => element.style.border = '2px solid red'", lod_date_element)
                print(f"lod_date: {lod_date} not matches")
    
    if lod_date == "" and match_response['Payout LOD Form Initial Upload'] == "not matched":  # require to handle where both are empt
        match_response['Payout LOD Form Recent Upload'] = "matched"
    page.screenshot(path="lod_date.png")

    # Locate the element containing the text "Letter of Direction"
    wait_and_click(page,XPATHS["history"])
    handle_angular_popups(page)
    wait_for_page_load(page)
    page.wait_for_timeout(glb_wait/2)
    # Safely handle empty or NaN values in the 'Payout LOD Form Recent Upload' column
    lod_payout_date = ""
    if "Payout LOD Form Recent Upload" in df.columns and not df["Payout LOD Form Recent Upload"].dropna().empty:
        lod_payout_date = df["Payout LOD Form Recent Upload"].dropna().iloc[0].strip()
    match_response['Payout LOD Form Recent Upload'] = "not matched" # explicitly set to not matched because validation will be in else statement
    page.wait_for_timeout(glb_wait/2)
    rows = page.query_selector_all("div.content-box div.content-title")
    count = 0
    for row in rows:
        lod_text = row.inner_text().strip() 
        if "Letter of Direction" in lod_text and "Completed" in lod_text:
            row.scroll_into_view_if_needed()
            page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", row)
            timestamp_element = row.query_selector("div.title-task")
            if timestamp_element:
                timestamp = timestamp_element.inner_text().strip()
                print(f"Extracted payout LOD: {timestamp}")
                if lod_payout_date.replace('\u00A0', ' ').strip() == timestamp.replace('\u00A0', ' ').strip():
                    match_response['Payout LOD Form Recent Upload'] = "matched"
                break
    if lod_payout_date == "" and match_response['Payout LOD Form Recent Upload'] == "not matched":  # require to handle where both are empt
        match_response['Payout LOD Form Recent Upload'] = "matched"
    page.screenshot(path="lod_payout_date.png")


    # Locate the element containing the text "Letter of Direction"
    wait_and_click(page,XPATHS["notifications"])
    handle_angular_popups(page)
    wait_for_page_load(page)
    page.wait_for_timeout(glb_wait/2)
    mmtg_num = ""
    if "CLASS Application #" in df.columns and not df["CLASS Application #"].dropna().empty:
        lod_payout_date = df["CLASS Application #"].dropna().iloc[0]

    match_response['CLASS Application #'] = "not matched" # explicitly set to not matched because validation will be in else statement
    
    rows = page.query_selector_all(XPATHS["mobile_notifications"])
    count = 0
    for row in rows:
        lod_text = row.inner_text().strip() 
        if str(mmtg_num).replace('\u00A0', ' ').strip() in lod_text.replace('\u00A0', ' ').strip():
            row.scroll_into_view_if_needed()
            row.click()
            page.wait_for_timeout(glb_wait/2)
            page.evaluate("element => { element.style.border = '2px solid green'; element.style.outline = 'none'; element.scrollIntoView({ behavior: 'smooth', block: 'center' }); }", row)
            match_response['CLASS Application #'] = "matched"
            break
    if mmtg_num == "" and match_response['CLASS Application #'] == "not matched":  # require to handle where both are empt
        match_response['CLASS Application #'] = "matched"
    page.screenshot(path="mmtg_num.png")

    print(f"**********************")
    print(f"Deal Summary Data: {match_response}")
    create_word_document(deal,match_response,df)

def open_deal(page,deal,glb_wait):
    wait_and_click(page,'(//span[@class="menuName"])[6]')
    wait_and_click(page,'(//span[@class="menuName"])[6]')
    wait_and_click(page,'.btn.clear-btn')
    wait_and_fill(page,'//input[@name="search"]', deal)
    wait_and_click(page,"button[id='searchButton']")
    
    wait_for_page_load(page)
    page.wait_for_timeout(glb_wait/2)
    key_value_pairs = {}
    deal_element = None
    try:
        deal_element = page.query_selector(f'(//a[contains(text(),"CIBC-2053438")])[1]')
        deal_element = page.query_selector(f'(//a[contains(text(),"CIBC-2053438")])[1]')
    except pwTimeoutError:
        key_value_pairs['Name'] = "Record not found"
        return
    page.query_selector(f'(//a[contains(text(),"CIBC-2053438")])[1]')
    
    wait_for_page_load(page)
    rc = wait_and_click_dont_throw_exception_if_element_not_exsist(page,'.applicationNo.clickable > a',expect_navigation=True)
    if rc == False:
        print(f"--- Not found {deal}")
        return
    click_if_exists(page,"#ConcurrentUserMsg > div > div > div.modal-footer.bdrtop > button.btn.btn-prospect-rev")
    click_if_exists(page,"//*[@id='CancelDealMsg']/div/div/div[3]/button")    

               
def mmtg_login(page, username, password,glb_wait):
    """Login to the portal using the provided username and password."""
    page.goto(login_url)
    wait_for_page_load(page)
    wait_and_fill(page, 'input[id="username"]', username)
    wait_and_fill(page, 'input[id="pwd"]', password)
    wait_for_page_load(page)
    wait_and_click(page, 'button[type="submit"]', expect_navigation=True)



def relogin_and_process(page, df, deal,glb_wait):
    """Relogin to the portal and process user IDs from the 'MA name' column."""

    # Iterate over the 'MA name' column in the DataFrame
    if "MA name" in df.columns:
        for user_id in df["MA name"].dropna():
            try:
                # Extract the number from the user_id field
                user_id_number = ''.join(filter(str.isdigit, str(user_id)))
                formatted_user_id = f"MTGQA{user_id_number}" if user_id_number else user_id
                
                print(f"Relogging in for user ID: {formatted_user_id}")
                mmtg_login(page, formatted_user_id, password,glb_wait)
                open_deal(page, deal,glb_wait)
                get_deal_summary_data(page,glb_wait,df,deal)

            except Exception as e:
                print(f"Error processing user ID {user_id}: {e}")
                print(traceback.format_exc())

def main_execution(deals_input):
    temp_file_path = None
    browser = None
    page = None
    workbook = None
    login_url, username, password, supportUser, supportPassword, glb_wait, headless_flag
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)  # Remove the default sheet
    # print(f"excel_file: {file_path}")
    # if len(file_path.strip())<3:
    #     print(f"***************************")
    #     print(f"Excel Outpul file is missing so exiting.")
    #     print(f"***************************")
    #     exit(0)
    # trace_dir = os.path.join(os.path.dirname(file_path), "ss_traces")
    # os.makedirs(trace_dir, exist_ok=True)
    # trace_file_path = os.path.join(trace_dir, f"trace_{int(time.time())}.zip")


    p = None
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=headless_flag)  # Set headless=True for headless mode    
            page = browser.new_page()
            mmtg_login(page, supportUser, password,glb_wait)



            # Process deals input into a list
            deals = [deal.strip() for deal in deals_input.split(",") if deal.strip()]
            print(f"Deals to process: {deals}")

            # ################################temp
            # df = pd.read_excel(r"C:\Users\YadavArv\Desktop\Arvind\Repo\python\mmtg_parser\downloads\10782 report for RESL credit applications_2025-04-09 06_37 AM.xlsx")
            # for deal in deals:
            #     if "CLASS Application #" in df.columns:
            #         matching_rows = df[df["CLASS Application #"] == deal]
            #         if not matching_rows.empty:
            #             print(f"Processing rows where deal matches: {deal}")
            #             for _, row in matching_rows.iterrows():
            #                 print(f"Processing row: {row.to_dict()}")
            #                 relogin_and_process(page, matching_rows, deal,glb_wait)  # Pass only the matched rows
                                            
            #         else:
            #             print(f"No matching rows found for deal: {deal}")

            # return

            # ################################ temp

            wait_for_page_load(page)
            wait_and_click(page,'(//span[@class="menuName"])[8]')
            wait_for_page_load(page)
            page.wait_for_timeout(glb_wait/2)
            wait_and_click(page,'//*[@id="quickLinkList"]/li[5]/a')
            wait_for_page_load(page)
            page.wait_for_timeout(glb_wait/2)
            record_found = False
            while True:
                if record_found:
                    print("Record found. Exiting loop after processing all rows.")
                    page.wait_for_timeout(glb_wait/2)
                    break  # Exit the loop after processing all rows
                max_wait_time = 120  # Maximum wait time in seconds
                elapsed_time = 0
                while not page.query_selector_all("div.row.assets-tbl-body"):
                    page.wait_for_timeout(glb_wait/2)
                    elapsed_time += 1
                    if elapsed_time >= max_wait_time:
                        print("Timeout: Rows not found within the maximum wait time.")
                        return
                rows = page.query_selector_all("div.row.assets-tbl-body")
                for row in rows:
                    try:
                        # Safely query each element and handle missing elements
                        report_type_element = row.query_selector("div:nth-of-type(1)")
                        report_type = report_type_element.inner_text().strip() if report_type_element else ""

                        to_date_element = row.query_selector("div:nth-of-type(3)")
                        to_date = to_date_element.inner_text().strip() if to_date_element else ""

                        report_time_element = row.query_selector("div:nth-of-type(5)")
                        report_time = report_time_element.inner_text().strip() if report_time_element else ""

                        report_name_element = row.query_selector("div:nth-of-type(6)")
                        report_name = report_name_element.inner_text().strip() if report_name_element else ""

                        status_element = row.query_selector("div:nth-of-type(7)")
                        status = status_element.inner_text().strip() if status_element else ""

                        print(f"Report Name: {report_name}, Status: {status}")

                        if "RESL" in report_name and status == "Completed":
                            download_icon = row.query_selector("img[id^='asset_TabelHed_del_icon']")
                            df = handle_file_download(page, download_icon)
                            # if df is not None:
                            #     print("Iterating over each column in the DataFrame:")
                            #     for column in df.columns:
                            #         print(f"Column: {column}")
                            #         print(df[column].head())  # Display the first few rows of the column
                            record_found = True
                            for deal in deals:
                                # Check if the deal matches any value in df["CLASS Application #"]
                                if "CLASS Application #" in df.columns:
                                    matching_rows = df[df["CLASS Application #"] == deal]
                                    if not matching_rows.empty:
                                        print(f"Processing rows where deal matches: {deal}")
                                        for _, row in matching_rows.iterrows():
                                            print(f"Processing row: {row.to_dict()}")
                                            relogin_and_process(page, matching_rows, deal,glb_wait)  # Pass only the matched rows
                                            
                                    else:
                                        print(f"No matching rows found for deal: {deal}")

                            break
                    except Exception as e:
                        print(f"Error processing row: {e}")

                # Add a delay before navigating to the next page to ensure all rows are processed
                page.wait_for_timeout(glb_wait/2)  # Adjust the delay as needed

                # Check if there is a next page button and click it
                next_page_button = page.query_selector("li.pagination-next.ng-star-inserted > a")
                if next_page_button and not next_page_button.is_disabled():
                    next_page_button.click()
                    wait_for_page_load(page)
                else:
                    print("No more pages to navigate.")
                    break


    except Exception as e:
        print(f"Error in run_script: {e}")
        print(traceback.format_exc())
    finally:

        if browser:
            try:
                print("Closing the browser...")
                p.stop()
            except Exception as e:
                print(f"Error while closing the browser: {e}")


main_execution("CIBC-2023127,CIBC-2053438,CIBC-2054810")

