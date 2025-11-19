from docx import Document
from docx.shared import Inches


def initialize_report(outputfilename, title_text):
    """
            This Function initializes Test Result Report document
    """
    document = Document()
    document.add_heading(title_text, 1)
    document.save(outputfilename)

def write_text(outputfilename, input_text):
    """
        This Function writes text to Test Results Report document
    """
    document = Document(outputfilename)
    document.add_paragraph()
    document.add_paragraph(input_text)
    document.save(outputfilename)

def write_heading(outputfilename, input_text):
    """
           This Function writes heading to Test Results Report document
    """
    document = Document(outputfilename)
    document.add_paragraph()
    document.add_heading(input_text, level=2)
    document.save(outputfilename)

def write_bullet(outputfilename, input_text):
    """
            This Function writes bullets to Test Results Report document
    """
    document = Document(outputfilename)
    document.add_paragraph()
    document.add_paragraph(input_text, style='List Bullet')
    document.save(outputfilename)

def attach_picture(outputfilename, pic_filename):
    """
           This Function attaches images to Test Results Report document
    """
    document = Document(outputfilename)
    document.add_paragraph()
    document.add_picture(pic_filename, width=Inches(6))
    document.save(outputfilename)
